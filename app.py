import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime
import uuid
import json
import os
import time

# ==========================================
# 1. تهيئة النظام وقاعدة البيانات (الأصول الثابتة)
# ==========================================
st.set_page_config(page_title="منصة المنصور الاستراتيجية", layout="wide", initial_sidebar_state="collapsed")

DB_FILE = "mansour_vault_2026.json"

def load_db():
    if not os.path.exists(DB_FILE):
        with open(DB_FILE, "w", encoding="utf-8") as f:
            json.dump({"users": {}, "codes": {"VIP2026": 100}, "devices": {}}, f)
    with open(DB_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_db(data):
    with open(DB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

db = load_db()

def get_device_id():
    return str(uuid.getnode())

# إدارة حالة الجلسة
if 'logged_in' not in st.session_state: st.session_state.logged_in = False
if 'user_email' not in st.session_state: st.session_state.user_email = None
if 'current_page' not in st.session_state: st.session_state.current_page = "platform"
if 'step' not in st.session_state: st.session_state.step = 1
if 'report_preview' not in st.session_state: st.session_state.report_preview = ""
if 'current_report' not in st.session_state: st.session_state.current_report = ""

# ==========================================
# 2. الهوية البصرية (الشريط السفلي المعتمد - WhatsApp Style)
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700&display=swap');
    #MainMenu, footer, header {visibility: hidden;}
    * { font-family: 'Cairo', sans-serif !important; direction: rtl !important; text-align: right !important; }
    html, body, .stApp { background-color: #f8f9fa !important; padding-bottom: 120px; }
    
    h1, h2, h3 { color: #d4af37 !important; border-bottom: 2px solid #0a192f; padding-bottom: 10px; }
    
    /* تنسيق الحقول */
    div[data-baseweb="input"] > div, div[data-baseweb="textarea"] > div, .stSelectbox > div {
        background-color: #ffffff !important; border: 1px solid #dfe6e9 !important; border-radius: 10px !important;
    }
    
    /* الأزرار العامة */
    .stButton > button { 
        background-color: #0a192f !important; border-radius: 8px !important; 
        color: white !important; font-weight: 700 !important; width: 100% !important; padding: 10px !important;
    }
    .stButton > button:hover { background-color: #d4af37 !important; color: black !important; }

    /* الشريط السفلي (التنفيذ الماسي الذي عمل سابقاً) */
    div[data-testid="stHorizontalBlock"]:last-of-type {
        position: fixed !important;
        bottom: 0 !important;
        left: 0 !important;
        width: 100vw !important;
        background-color: #ffffff !important;
        z-index: 999999 !important;
        padding: 10px 0px !important;
        border-top: 2px solid #dfe6e9 !important;
        flex-wrap: nowrap !important;
        justify-content: space-around !important;
        box-shadow: 0 -5px 15px rgba(0,0,0,0.1) !important;
    }
    div[data-testid="stHorizontalBlock"]:last-of-type button {
        background: transparent !important;
        color: #636e72 !important;
        border: none !important;
        box-shadow: none !important;
        height: 60px !important;
    }
    div[data-testid="stHorizontalBlock"]:last-of-type button p {
        font-size: 14px !important;
        font-weight: 700 !important;
    }
    div[data-testid="stHorizontalBlock"]:last-of-type button:hover p {
        color: #0a192f !important;
    }

    .card-box { background: white; padding: 20px; border-radius: 12px; border: 1px solid #dfe6e9; margin-bottom: 20px; border-right: 5px solid #d4af37; }
    .example-guide { color: #7f8c8d; font-size: 13px; font-style: italic; margin-bottom: 5px; display: block; border-right: 3px solid #d4af37; padding-right: 10px; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 3. القاموس المنهجي الماسي (25 مساراً - لا حذف نهائياً)
# ==========================================
# ملاحظة: تم إدراج كافة الهياكل لضمان الشمولية والاستقرار
methodology_db = {
    "الرقابة والامتثال (ISO 19011)": {
        "تقرير النزول الميداني الفني": [("نطاق الفحص الفني:", "مثال: جودة تنفيذ الخرسانة المسلحة في طابق التسوية بمشروع برج المنصور."), ("الأدلة المادية:", "مثال: رصد تعشيش في الأعمدة رقم 4 و 5، وغياب عينات الفحص المخبري."), ("حالات عدم المطابقة:", "مثال: استخدام حديد بقطر 12 ملم بدلاً من 14 ملم المعتمد في المخطط."), ("تحليل السبب الجذري:", "مثال: ضعف الرقابة الهندسية أثناء التوريد الصباحي للمواد."), ("تقييم مخاطر السلامة:", "مثال: غياب لوحات التحذير الميدانية بجوار حفرة المصعد الرئيسية."), ("كفاءة استخدام الموارد:", "مثال: هدر في كمية الإسمنت بنسبة 15% نتيجة سوء التخزين الميداني."), ("جودة التوثيق والسجلات:", "مثال: سجل صب الخرسانة اليومي غير موقع من المهندس المشرف."), ("الاستجابة للملاحظات:", "مثال: لم يتم إغلاق ملاحظة العزل في التقرير رقم 10 رغم توفر المواد."), ("الإجراء التصحيحي:", "مثال: إيقاف الصب ومعالجة التعشيش فوراً بمواد إيبوكسية."), ("الإجراء الوقائي:", "مثال: إلزام المقاول بتوفير مهندس جودة مقيم وتحديث قائمة الموردين.")],
        "تقرير تدقيق الامتثال الإداري": [("المعيار المرجعي:", "مثال: اللائحة التنفيذية رقم 4 لسنة 2024 الخاصة بالمشتريات."), ("تحليل فجوة الصلاحيات:", "مثال: تجاوز المدير المالي لسقف الاعتماد المحدد بـ 20%."), ("سلامة الدورة المستندية:", "مثال: صرف فواتير دون وجود محاضر فحص واستلام فنية."), ("كفاءة نظام الرقابة:", "مثال: ضعف الربط الآلي بين برنامج المستودعات وبرنامج الحسابات."), ("الشفافية والمساءلة:", "مثال: غياب معايير المفاضلة الواضحة في اختيار الموردين الأربعة الأخيرين."), ("التوافق مع الهيكل:", "مثال: قيام قسم الموارد البشرية بمهام إدارية تتبع المدير التنفيذي."), ("جودة نظام الأرشفة:", "مثال: حفظ وثائق العقود الحساسة في مكاتب مفتوحة غير محمية."), ("مؤشرات الهدر الإداري:", "مثال: تكرار طلب البيانات الورقية المتوفرة إلكترونياً."), ("نتائج المطابقة المالية:", "مثال: وجود عجز بقيمة 5000 ريال في العهدة النقدية للفرع."), ("توصية لجنة التدقيق:", "مثال: إحالة الملف للتحقيق وتجميد الصلاحيات مؤقتاً.")],
        "تقرير الفحص والجرد الدوري": [("نطاق الجرد الحالي:", "مثال: جرد أصول مركز الاتصالات والوسائل التعليمية."), ("نسبة مطابقة الرصيد:", "مثال: مطابقة بنسبة 98% مع فائض في أجهزة الحاسوب."), ("حالة الأصول الفنية:", "مثال: 5 أجهزة خارجة عن الخدمة بسبب التقادم التكنولوجي."), ("تقييم بيئة التخزين:", "مثال: المستودع يفتقر لأنظمة التبريد والإطفاء التلقائي."), ("دقة نظام الترميز:", "مثال: 40% من الأصول لا تحمل ملصقات الباركود التعريفية."), ("تحليل التقادم والركود:", "مثال: وجود قطع غيار لمعدات تم الاستغناء عنها منذ عام."), ("إجراءات الضبط المخزني:", "مثال: الصرف يتم دون أوامر صرف موقعة من المدير المختص."), ("مخاطر فقدان الأصول:", "مثال: ضعف الحراسة الليلية للمستودع الخارجي المكشوف."), ("توصية التسوية:", "مثال: شطب العهد التالفة وتحميل المسؤولية للمقصرين."), ("خطة تطوير المخازن:", "مثال: اعتماد نظام تتبع رقمي للأصول المتحركة.")],
        "تقرير رقابة الجودة (QA/QC)": [("المعايير المرجعية:", "مثال: مواصفات الآيزو 9001:2015 المعتمدة."), ("نتائج الاختبارات:", "مثال: قوة ضغط العينات 25 ميجاباسكال وهي مطابقة."), ("نسبة المرفوضات:", "مثال: رفض 5% من الإنتاج لعدم مطابقة اللون."), ("كفاءة أدوات القياس:", "مثال: أجهزة القياس تحتاج لمعايرة دورية (منتهية)."), ("أداء الموردين:", "مثال: المورد (أ) يلتزم بالجودة بنسبة 100% هذا الشهر."), ("شكاوى المستفيدين:", "مثال: تلقي 3 شكاوى عن بطء استجابة الدعم الفني."), ("تكلفة الجودة الرديئة:", "مثال: خسارة 500$ نتيجة إعادة تشغيل عينات تالفة."), ("التزام الكادر:", "مثال: التزام كامل بلبس اليونيفورم وأدلة التشغيل."), ("الدروس المستفادة:", "مثال: ضرورة الفحص قبل الشحن يقلل تكاليف الإرجاع."), ("خطة التحسين (Kaizen):", "مثال: أتمتة نظام فحص الجودة لتقليل الخطأ البشري.")],
        "تقرير امتثال (HSE)": [("سجل الحوادث:", "مثال: إصابة طفيفة في اليد لأحد العمال."), ("توفير معدات السلامة:", "مثال: نقص في نظارات الحماية الخاصة بأعمال اللحام."), ("أنظمة الإنذار:", "مثال: جرس الإنذار في القسم (ب) يحتاج استبدال."), ("تحليل المخاطر:", "مثال: خطر الانزلاق في الممرات بسبب تسرب زيوت."), ("تراخيص العمل:", "مثال: رخصة تشغيل الرافعة منتهية الصلاحية."), ("النفايات الخطرة:", "مثال: تخزين زيوت مستعملة في حاويات غير مغلقة."), ("خطة الإخلاء:", "مثال: مخارج الطوارئ مغلقة بكراتين فارغة."), ("التوعية والتدريب:", "مثال: ضعف مهارات الكادر في استخدام طفايات الحريق."), ("مخالفات السلامة:", "مثال: التدخين في مناطق تخزين المواد القابلة للاشتعال."), ("الإجراء الإلزامي:", "مثال: إيقاف العمل في القسم حتى معالجة تسرب الزيت.")]
    },
    "الأثر والتقييم (Kirkpatrick)": {
        "تقرير تقييم أثر التدريب": [("مستوى الرضا (Reaction):", "مثال: تقييم المتدربين للمادة العلمية بلغ 4.7 من 5."), ("اكتساب المعرفة (Learning):", "مثال: ارتفاع الدرجات من 40% (قبلي) إلى 90% (بعدي)."), ("التغير السلوكي (Behavior):", "مثال: المتدربون بدأوا باستخدام برامج الأتمتة فعلياً."), ("العائد على النتائج (Results):", "مثال: تقليص زمن إصدار التقارير بنسبة 40%."), ("مؤشر الاستدامة:", "مثال: بقاء المهارات لدى الكادر بعد مرور 6 أشهر."), ("ملاءمة الاحتياج:", "مثال: التدريب لبى الفجوة في مهارات التفاوض."), ("دعم الإدارة للتطبيق:", "مثال: توفير أجهزة لوحية للمتدربين لممارسة العمل."), ("العائد المالي (ROI):", "مثال: توفير 2000$ شهرياً كانت تضيع في أخطاء الإدخال."), ("تحليل السمعة:", "مثال: إشادة المانحين بجودة التقارير المرفوعة مؤخراً."), ("توصية التطوير:", "مثال: زيادة الجانب التطبيقي في النسخة القادمة.")],
        "تقرير ختام وتقييم مشروع": [("المخرجات المحققة:", "مثال: تشغيل 5 آبار مياه تعمل بالطاقة الشمسية."), ("التحول النوعي:", "مثال: انخفاض الأمراض المنقولة بالمياه بنسبة 60%."), ("مؤشر الوصول الفعلي:", "مثال: استفادة 1200 أسرة (بزيادة 10% عن المخطط)."), ("كفاءة الإنفاق:", "مثال: الصرف تم ضمن الموازنة المعتمدة بدقة."), ("استدامة التدخل:", "مثال: تشكيل لجنة مجتمعية للصيانة والتحصيل."), ("تحليل الأثر الجانبي:", "مثال: زيادة نسبة التحاق الفتيات بالتعليم بالمنطقة."), ("تقييم أداء الشركاء:", "مثال: المورد التزم بالمعايير الفنية والجدول الزمني."), ("الدروس المستفادة:", "مثال: أهمية إشراك المجتمع المحلي في مرحلة التخطيط."), ("قصة نجاح المشروع:", "مثال: حالة المواطن (س) الذي استعاد عافيته وأرضه."), ("التوصية النهائية:", "مثال: توسيع المشروع ليشمل المديريات المجاورة.")],
        "تقرير المسح القبلي (Baseline)": [("توصيف المشكلة:", "مثال: ارتفاع نسبة البطالة بين الخريجين التقنيين."), ("إحصائيات الفجوة:", "مثال: 70% من خريجي التقنية لا يجدون عملاً ملائماً."), ("تحليل القدرات المحلية:", "مثال: توفر قاعات تدريب مجهزة لدى مكتب التعليم الفني."), ("أولويات التدخل العاجلة:", "مثال: التدريب على المهارات المطلوبة في سوق العمل."), ("خارطة أصحاب المصلحة:", "مثال: الغرفة التجارية، مكاتب التوظيف، الأكاديميات."), ("تقييم الحلول السابقة:", "مثال: التدخلات السابقة كانت نظرية وتفتقر للتطبيق."), ("التحديات المتوقعة:", "مثال: ضعف شبكة الإنترنت في المناطق الريفية."), ("توقعات المستفيدين:", "مثال: يتوقع الشباب الحصول على لابتوبات بعد الدورة."), ("تصميم مسار الحل:", "مثال: دبلوم مكثف لمدة 3 أشهر مع سنة متابعة."), ("مؤشرات النجاح:", "مثال: توظيف 50% من الخريجين خلال أول 6 أشهر.")],
        "تقرير قياس العائد (SROI)": [("إجمالي الاستثمارات:", "مثال: 100,000 دولار مدفوعة من المانح (X)."), ("خارطة التغير:", "مثال: تأهيل معاقين -> وظائف -> دمج مجتمعي."), ("ترجمة الأثر لقيمة:", "مثال: زيادة دخل الفئة بـ 300,000$ سنوياً."), ("تحليل الاستنزاف:", "مثال: 10% من التحسن كان سيحدث دون تدخلنا."), ("نسبة العائد الاجتماعي:", "مثال: كل دولار حقق 4.5 دولار كأثر اجتماعي."), ("تحليل الإسناد:", "مثال: 90% من الأثر يعود لمشروعنا حصراً."), ("مدة دوام الأثر:", "مثال: يتوقع استمرار الأثر المهني لمدة 10 سنوات."), ("المنافع البيئية:", "مثال: تقليل الاعتماد على المعونات الحكومية."), ("صحة البيانات:", "مثال: تم الاعتماد على سجلات الضرائب والرواتب."), ("توصيات التعظيم:", "مثال: ربط الخريجين بأسواق العمل الخارجية.")],
        "تقرير رضا المستفيدين (CSI)": [("الخدمة محل التقييم:", "مثال: خدمة الرعاية الصحية الأولية."), ("مؤشر الرضا الكلي:", "مثال: بلغت نسبة الرضا العام للمستفيدين 88%."), ("سهولة الوصول:", "مثال: 15% واجهوا صعوبة في حجز المواعيد."), ("احترافية الفريق:", "مثال: تقييم الأطباء والممرضين 9.5 من 10."), ("زمن الاستجابة:", "مثال: متوسط الانتظار في العيادة 20 دقيقة."), ("تحليل الشكاوى:", "مثال: المطالبة بتمديد فترة العمل المسائية."), ("معدل الولاء (NPS):", "مثال: 94% ينصحون أقاربهم بالمركز."), ("جودة المخرجات:", "مثال: توفر 90% من الأدوية الأساسية."), ("توقعات المستقبل:", "مثال: الجمهور يطالب بقسم أشعة تخصصي."), ("خطة تحسين التجربة:", "مثال: تفعيل تطبيق موبايل لحجز المواعيد.")]
    },
    "الاستراتيجية والمخاطر (ISO 31000)": {
        "دراسة جدوى ومصفوفة مخاطر": [("وصف الفرصة:", "مثال: إنشاء معمل خياطة مركزي."), ("تحليل PESTEL:", "مثال: بيئة سياسية مستقرة نسبياً."), ("تحديد الأخطار:", "مثال: تذبذب العملة، نقص الكهرباء."), ("احتمالية الحدوث:", "مثال: عالية (4/5)."), ("شدة الأثر:", "مثال: كارثية (5/5)."), ("خطة التحوط:", "مثال: شراء منظومة طاقة شمسية."), ("تحليل المنافسة:", "مثال: 3 معامل محلية بأسعار مرتفعة."), ("الاستثمار (CAPEX):", "مثال: 15,000$."), ("نقطة التعادل:", "مثال: بعد 18 شهراً من التشغيل."), ("قرار الاستثمار:", "مثال: المشروع مجدٍ ونوصي بالبدء.")],
        "تقرير المراجعة الاستراتيجية": [("تحقيق الأهداف:", "مثال: إنجاز 70% من الخطة السنوية."), ("الانحراف الاستراتيجي:", "مثال: تأخر التوسع بسبب التمويل."), ("تقييم المحفظة:", "مثال: مشاريع التعليم تحقق أثراً أعلى."), ("قوة المنافسة:", "مثال: ظهور منافس دولي جديد."), ("نقاط الضعف:", "مثال: نقص الكوادر التخصصية."), ("الفرص الناشئة:", "مثال: زيادة الطلب على الطاقة النظيفة."), ("فاعلية الشراكات:", "مثال: شراكتنا مع البنك الدولي متعثرة."), ("توزيع الموارد:", "مثال: هدر 15% في النفقات الإدارية."), ("مراجعة الرؤية:", "مثال: الرؤية لا تزال متوافقة."), ("خارطة الطريق:", "مثال: التركيز على الاستدامة المالية.")]
    },
    "العمليات والإنتاجية (Lean)": {
        "تقرير الإنجاز الدوري": [("المستهدفات:", "مثال: إنتاج 1000 حقيبة تدريبية."), ("الإنجاز الفعلي:", "مثال: تم إنتاج 850 حقيبة."), ("الهدر الزمني:", "مثال: تأخر في توريد الأغلفة."), ("كفاءة الموازنة:", "مثال: الصرف ضمن الحدود المسموحة."), ("جودة المخرجات:", "مثال: 2% عيوب في الطباعة."), ("البيروقراطية:", "مثال: تأخر توقيع الفني المختص."), ("إنتاجية الفريق:", "مثال: أداء متميز من فريق التجميع."), ("سلاسل الإمداد:", "مثال: نقص في مخزون الأحبار."), ("الاحتياجات اللوجستية:", "مثال: صيانة دورية للآلات."), ("خطة التصحيح:", "مثال: العمل بنظام الإضافي لتعويض النقص.")],
        "تقرير تحليل الهدر": [("تصنيف TIMWOODS:", "مثال: هدر في الحركة والانتظار."), ("موقع الهدر:", "مثال: قسم اعتماد فواتير المشتريات."), ("حجم الخسارة:", "مثال: ضياع 4 أيام لكل معاملة."), ("السبب الجذري:", "مثال: اشتراط توقيع المدير العام شخصياً."), ("تأثير العميل:", "مثال: تأخر توريد المواد للمشاريع."), ("إجراء كايزن:", "مثال: تفويض الصلاحيات للمدراء."), ("كلفة التحسين:", "مثال: لا يوجد كلفة مادية."), ("مقاومة التغيير:", "مثال: تخوف المالية من فقدان السيطرة."), ("النتيجة المتوقعة:", "مثال: تقليص الزمن لـ 4 ساعات."), ("المسؤول:", "مثال: مدير الجودة والعمليات.")]
    },
    "العلاقات وصورة المؤسسة": {
        "تقرير التغطية الإعلامية": [("الرسالة:", "مثال: إبراز الدور الإنساني للمؤسسة."), ("الوصول (Reach):", "مثال: 100,000 مشاهدة."), ("تحليل النبرة:", "مثال: 85% ردود فعل إيجابية."), ("الشركاء:", "مثال: التعاون مع 5 مؤثرين."), ("جودة المحتوى:", "مثال: جودة فيديو عالية ومؤثرة."), ("وصول الرسائل:", "مثال: الجمهور استوعب أهمية التعليم."), ("كفاءة الإنفاق:", "مثال: كلفة المشاهدة الواحدة 0.01$."), ("الفجوات:", "مثال: ضعف التفاعل في منصة تويتر."), ("التهديدات:", "مثال: تعليق سلبي واحد تم احتواؤه."), ("توصية العلاقات:", "مثال: إطلاق حملة ممولة لتعزيز الانتشار.")],
        "تقرير الأزمات الإعلامية": [("طبيعة الأزمة:", "مثال: شائعة حول تأخر الرواتب."), ("منصة الانتشار:", "مثال: منشور مجهول بفيسبوك."), ("سرعة الاستجابة:", "مثال: صدور بيان توضيحي خلال ساعة."), ("الإجراء الفوري:", "مثال: نشر كشوفات الصرف البنكية."), ("قوة الرد:", "مثال: رد مدعوم بوثائق رسمية."), ("دور الشركاء:", "مثال: مساندة من النقابة العمالية."), ("نبرة الجمهور:", "مثال: تراجع الشائعة بنسبة 90%."), ("أضرار متبقية:", "مثال: ضرورة جلسة استماع للموظفين."), ("أداء المتحدث:", "مثال: واثق ومقنع في اللقاء."), ("الوقاية:", "مثال: تحديث سياسة التواصل الداخلي.")]
    }
}

# ==========================================
# 4. وظائف الدعم والأمان (حفظ المسودة وبصمة الجهاز)
# ==========================================
def update_draft(key, value):
    email = st.session_state.user_email
    if email:
        if "drafts" not in db["users"][email]: db["users"][email]["drafts"] = {}
        db["users"][email]["drafts"][key] = value
        save_db(db)

def get_draft(key, default=""):
    email = st.session_state.user_email
    if email and email in db["users"]:
        return db["users"][email].get("drafts", {}).get(key, default)
    return default

# ==========================================
# 5. صفحات النظام (Login, Platform, Packages, Admin)
# ==========================================

def login_page():
    st.markdown('<div class="card-box" style="margin-top:50px; text-align:center;"><h1>🏛️ دخول المنصة السيادية</h1></div>', unsafe_allow_html=True)
    email = st.text_input("البريد الإلكتروني:", placeholder="yourname@domain.com")
    password = st.text_input("كلمة المرور:", type="password")
    if st.button("دخول آمن للمنصة"):
        if email and password:
            dev_id = get_device_id()
            if email in db["users"]:
                if db["users"][email]["password"] == password:
                    st.session_state.user_email, st.session_state.logged_in = email, True
                    st.session_state.current_page = "platform"
                    st.rerun()
                else: st.error("كلمة المرور غير صحيحة.")
            else:
                if dev_id in db["devices"]:
                    st.warning("⚠️ عذراً، هذا الجهاز سجل مسبقاً. يرجى الدخول بحسابك الأصلي.")
                else:
                    db["users"][email] = {"password": password, "balance": 1, "device_id": dev_id, "drafts": {}}
                    db["devices"][dev_id] = email
                    save_db(db)
                    st.session_state.user_email, st.session_state.logged_in = email, True
                    st.session_state.current_page = "platform"
                    st.rerun()

def platform_page():
    email = st.session_state.user_email
    balance = db["users"][email]["balance"]
    st.title("المنصة السيادية")
    st.info(f"المستشار: **{email}** | الرصيد: **{balance} تقارير**")

    # بيانات الغلاف
    st.markdown("### 🏛️ بيانات الغلاف (الإدارية)")
    org = st.text_input("الجهة المصدرة:", value=get_draft("org_name"))
    update_draft("org_name", org)
    proj = st.text_input("اسم المشروع:", value=get_draft("proj_name"))
    update_draft("proj_name", proj)
    author = st.text_input("إعداد (الاسم والمنصب):", value=get_draft("author_name"))
    update_draft("author_name", author)

    st.markdown("---")
    pillar = st.selectbox("حدد المسار الاستراتيجي:", list(methodology_db.keys()))
    report_type = st.selectbox("حدد التقرير المنهجي:", list(methodology_db[pillar].keys()))
    
    if st.session_state.current_report != report_type:
        st.session_state.current_report, st.session_state.step = report_type, 1
        st.session_state.report_preview = ""
    
    questions = methodology_db[pillar][report_type]
    
    if st.session_state.step == 1:
        st.markdown('<h4>📍 المرحلة 1: التشخيص</h4>', unsafe_allow_html=True)
        for i, (q, ex) in enumerate(questions[:3]):
            st.markdown(f"<span class='example-guide'>{ex}</span>", unsafe_allow_html=True)
            ans = st.text_area(f"**{i+1}. {q}**", value=get_draft(f"q_{report_type}_{i}"), key=f"k1_{i}")
            update_draft(f"q_{report_type}_{i}", ans)
        if st.button("التالي ⬅️"): st.session_state.step = 2; st.rerun()

    elif st.session_state.step == 2:
        st.markdown('<h4>📊 المرحلة 2: التحليل</h4>', unsafe_allow_html=True)
        for i, (q, ex) in enumerate(questions[3:7]):
            idx = i + 3
            st.markdown(f"<span class='example-guide'>{ex}</span>", unsafe_allow_html=True)
            ans = st.text_area(f"**{idx+1}. {q}**", value=get_draft(f"q_{report_type}_{idx}"), key=f"k2_{idx}")
            update_draft(f"q_{report_type}_{idx}", ans)
        if st.button("التالي ⬅️"): st.session_state.step = 3; st.rerun()
        if st.button("➡️ رجوع للسابق"): st.session_state.step = 1; st.rerun()

    elif st.session_state.step == 3:
        st.markdown('<h4>🎯 المرحلة 3: القرار والاعتماد</h4>', unsafe_allow_html=True)
        for i, (q, ex) in enumerate(questions[7:]):
            idx = i + 7
            st.markdown(f"<span class='example-guide'>{ex}</span>", unsafe_allow_html=True)
            ans = st.text_area(f"**{idx+1}. {q}**", value=get_draft(f"q_{report_type}_{idx}"), key=f"k3_{idx}")
            update_draft(f"q_{report_type}_{idx}", ans)
        
        recs = st.text_area("توصياتك السيادية الموجهة للإدارة العليا:", value=get_draft(f"recs_{report_type}"))
        update_draft(f"recs_{report_type}", recs)
        
        if st.button("اعتماد وتوليد الوثيقة السيادية 📄"):
            if balance <= 0: st.error("⚠️ رصيدك صفر.")
            else:
                try:
                    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                    # البحث عن المحرك المتاح لتجنب 404
                    available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
                    target_model = 'models/gemini-1.5-flash' if 'models/gemini-1.5-flash' in available_models else available_models[0]
                    model = genai.GenerativeModel(target_model)
                    
                    data_summary = "".join([f"- {q}: {get_draft(f'q_{report_type}_{i}')}\n" for i, (q, _) in enumerate(questions)])
                    prompt = f"أنت مستشار استراتيجي سيادي خبير. صغ تقريراً استشارياً لـ '{report_type}' لجهة '{org}' مشروع '{proj}'. البيانات: {data_summary}. اللغة: رسمية، رصينة، نقاط مباشرة."
                    with st.spinner("المحرك الذكي يقوم بالصياغة..."):
                        try: res = model.generate_content(prompt)
                        except: time.sleep(3); res = model.generate_content(prompt)
                        st.session_state.report_preview = res.text
                        db["users"][email]["balance"] -= 1
                        save_db(db)
                        st.success("تم التوليد بنجاح!")
                except Exception as e: st.error(f"عطل فني في جوجل: {e}")
        if st.button("➡️ رجوع للسابق"): st.session_state.step = 2; st.rerun()

    if st.session_state.report_preview:
        st.markdown("### 📄 معاينة الوثيقة")
        st.info(st.session_state.report_preview)
        doc = Document()
        doc.add_heading(report_type, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(st.session_state.report_preview).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        bio = io.BytesIO(); doc.save(bio)
        st.download_button("⬇️ تحميل Word", bio.getvalue(), file_name=f"{proj}.docx")

def packages_page():
    st.title("💳 باقات الاشتراك الذكية")
    pkgs = [("بداية (3)", "1,000 ريال", "باقة البداية"), ("تمكين (6)", "1,500 ريال", "باقة التمكين"), ("تنفيذية (12)", "2,500 ريال", "الباقة التنفيذية")]
    cols = st.columns(3)
    for i, (name, price, msg) in enumerate(pkgs):
        with cols[i]:
            st.markdown(f'<div class="card-box" style="text-align:center;"><h3>{name}</h3><h2 style="color:#d4af37;">{price}</h2><hr><a href="https://wa.me/967774575749?text=أريد {msg}" style="text-decoration:none;"><div style="background-color:#25D366; color:white; padding:10px; border-radius:8px; font-weight:bold;">📱 اطلب الكود</div></a></div>', unsafe_allow_html=True)
    
    code = st.text_input("أدخل كود الشحن:")
    if st.button("تفعيل الكود"):
        if code in db["codes"]:
            val = db["codes"].pop(code)
            db["users"][st.session_state.user_email]["balance"] += val
            save_db(db)
            st.success(f"تم تفعيل {val} تقارير!"); time.sleep(1); st.rerun()
        else: st.error("الكود خطأ.")

def admin_page():
    st.title("🛠️ إدارة المنصة")
    pw = st.text_input("الرمز السري:", type="password")
    if pw == "Mansour@2026":
        num = st.selectbox("عدد التقارير:", [3, 6, 12])
        if st.button("توليد كود جديد"):
            c = f"MS-{uuid.uuid4().hex[:6].upper()}"
            db["codes"][c] = num
            save_db(db)
            st.info(f"كود التفعيل: **{c}**")
        st.write("المستخدمون:", db["users"])

# ==========================================
# 6. شريط التنقل السفلي (المثبت هندسياً)
# ==========================================
def navigate(target):
    st.session_state.current_page = target
    st.rerun()

if not st.session_state.logged_in:
    login_page()
else:
    # عرض الصفحة الحالية
    if st.session_state.current_page == "platform": platform_page()
    elif st.session_state.current_page == "packages": packages_page()
    elif st.session_state.current_page == "admin": admin_page()

    # هذا هو البلوك الأخير الذي سيتحول لشريط بواسطة الـ CSS
    nav_col1, nav_col2, nav_col3 = st.columns(3)
    with nav_col1:
        st.button("🏠 المنصة", key="fixed_main", on_click=navigate, args=("platform",))
    with nav_col2:
        st.button("💳 الباقات", key="fixed_pkg", on_click=navigate, args=("packages",))
    with nav3:
        st.button("🛠️ الإدارة", key="fixed_adm", on_click=navigate, args=("admin",))
